#!/usr/bin/env python3
"""
Montaj Sahası Günlük Rapor Botu  —  OpenAI (Whisper + GPT-4o)
==============================================================
Taslak Word dosyasını kopyalar, içine veri yazar.

KURULUM:
  pip install python-telegram-bot openai python-docx pillow lxml

ÇALIŞTIRMA:
  python montaj_rapor_bot.py
"""

import os, logging, tempfile, json, shutil, copy, re
from datetime import datetime
from pathlib import Path
from io import BytesIO

from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import (
    Application, CommandHandler, MessageHandler,
    CallbackQueryHandler, ContextTypes, filters
)
import openai
from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ══════════════════════════════════════════════
#  YAPILANDIRMA
# ══════════════════════════════════════════════
TELEGRAM_TOKEN = os.getenv("TELEGRAM_TOKEN")
OPENAI_KEY     = os.getenv("OPENAI_KEY")   # sk-... ile başlar
TASLAK_PATH    = "MONTAJ_RAPORU_TASLAK.docx"  # bot ile aynı klasörde olmalı
OUTPUT_DIR     = Path("raporlar")
OUTPUT_DIR.mkdir(exist_ok=True)

# ══════════════════════════════════════════════
logging.basicConfig(format="%(asctime)s [%(levelname)s] %(message)s", level=logging.INFO)
log = logging.getLogger(__name__)
openai_client = openai.AsyncOpenAI(api_key=OPENAI_KEY)

GUNLER = {
    "Monday":"Pazartesi","Tuesday":"Salı","Wednesday":"Çarşamba",
    "Thursday":"Perşembe","Friday":"Cuma","Saturday":"Cumartesi","Sunday":"Pazar"
}
def gun_adi(tarih_str: str) -> str:
    try:
        dt = datetime.strptime(tarih_str, "%d.%m.%Y")
        return GUNLER.get(dt.strftime("%A"), "")
    except:
        return ""

def bugun() -> str:
    return datetime.now().strftime("%d.%m.%Y")


# ──────────────────────────────────────────────
#  OTURUM
# ──────────────────────────────────────────────
def yeni_oturum() -> dict:
    return {
        "tarih": datetime.now().strftime("%d/%m/%Y"),
        "is_emri": "", "firma": "", "ulke_sehir": "", "satis": "",
        "foremen": "", "personel": [], "baslangic": "", "bitis": "",
        "gerceklesen": "", "planlanan": "", "tesis_tipi": "", "notlar": "",
        "kontrol_eden": "", "onaylayan": "",
        "gunluk":  [],   # [{"tarih":"28.11.2025","gun":"Cuma","maddeler":[],"saat":"","fazla":""}]
        "aksalik": [],   # [{"aciklama":"...","fotograflar":[bytes,...]}]
        "musteri": [],
        "oneriler": [],
        "_mod": None,
        "_pending_fmt": None,
        "_pending_text": "",
        "_pending_photos": [],
    }

sessions: dict[int, dict] = {}
def get_session(uid: int) -> dict:
    if uid not in sessions:
        sessions[uid] = yeni_oturum()
    return sessions[uid]


# ──────────────────────────────────────────────
#  OPENAI
# ──────────────────────────────────────────────
async def transcribe(path: str) -> str:
    with open(path, "rb") as f:
        r = await openai_client.audio.transcriptions.create(
            model="whisper-1", file=f, language="tr")
    return r.text

SISTEM_GUNLUK = f"""Sen montaj sahası rapor asistanısın. Bugünün tarihi: {bugun()}.
Kullanıcının anlattığı günü JSON olarak döndür. SADECE JSON yaz, başka hiçbir şey.
{{
  "tarih": "GG.AA.YYYY",
  "maddeler": ["madde 1", "madde 2"],
  "calisma_saati": "08:00-18:00",
  "fazla_mesai": ""
}}
Tarih söylenmemişse BUGÜNÜN tarihini yaz: {bugun()}
Saatler söylenmemişse "08:00-18:00" yaz. Her işi ayrı madde yap."""

SISTEM_AKSALIK = "Montaj sahasındaki aksaklığı kısa ve net Türkçe teknik cümle olarak yaz. SADECE açıklama, başka hiçbir şey."
SISTEM_LISTE   = "Metni kısa Türkçe maddelere çevir. Her madde ayrı satırda. Numara veya tire koyma."

SISTEM_TOPLU = """Sen deneyimli bir montaj sahası rapor asistanısın.
Kullanıcı sana bir tarih aralığı ve o süreçte yapılan işlerin genel bir özetini verecek.
Bu özeti, tarih aralığındaki her gün için ayrı günlük kayıtlara böl.
Montaj projelerinde tipik iş akışını (hazırlık → yapısal montaj → elektrik → test → devreye alma) göz önünde bulundur.
Hangi işin hangi güne düşebileceğini mantıklı şekilde dağıt. Hafta sonları da çalışıldığını varsay.

SADECE JSON döndür, başka hiçbir şey yazma:
[
  {
    "tarih": "GG.AA.YYYY",
    "maddeler": ["madde 1", "madde 2"],
    "calisma_saati": "08:00-18:00",
    "fazla_mesai": ""
  }
]"""

async def gpt_gunluk(text: str) -> dict:
    r = await openai_client.chat.completions.create(
        model="gpt-4o", max_tokens=600,
        messages=[{"role":"system","content":SISTEM_GUNLUK},
                  {"role":"user","content":text}])
    raw = r.choices[0].message.content.strip().replace("```json","").replace("```","").strip()
    data = json.loads(raw)
    # Tarih boşsa bugünü yaz
    if not data.get("tarih"):
        data["tarih"] = bugun()
    data["gun"] = gun_adi(data["tarih"])
    return data

async def gpt_aksalik(text: str) -> str:
    r = await openai_client.chat.completions.create(
        model="gpt-4o", max_tokens=300,
        messages=[{"role":"system","content":SISTEM_AKSALIK},
                  {"role":"user","content":text}])
    return r.choices[0].message.content.strip()

async def gpt_liste(text: str) -> list:
    r = await openai_client.chat.completions.create(
        model="gpt-4o", max_tokens=300,
        messages=[{"role":"system","content":SISTEM_LISTE},
                  {"role":"user","content":text}])
    return [l.strip() for l in r.choices[0].message.content.strip().splitlines() if l.strip()]

async def gpt_toplu(tarih_aralik: str, paragraf: str) -> list:
    """Toplu paragrafı tarih aralığına göre gün gün böler."""
    prompt = f"Tarih aralığı: {tarih_aralik}\n\nYapılan işler (genel özet):\n{paragraf}"
    r = await openai_client.chat.completions.create(
        model="gpt-4o", max_tokens=2000,
        messages=[{"role":"system","content":SISTEM_TOPLU},
                  {"role":"user","content":prompt}])
    raw = r.choices[0].message.content.strip().replace("```json","").replace("```","").strip()
    gunler = json.loads(raw)
    for g in gunler:
        if not g.get("tarih"): continue
        g["gun"] = gun_adi(g["tarih"])
    return gunler


# ──────────────────────────────────────────────
#  XML YARDIMCILARI
# ──────────────────────────────────────────────
NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
W  = "{%s}" % NS

def yeni_run(text: str, bold=False, italic=False, size=22) -> etree._Element:
    """Şablonun font stiline uygun <w:r> oluşturur."""
    r = etree.Element(W+"r")
    rpr = etree.SubElement(r, W+"rPr")
    fonts = etree.SubElement(rpr, W+"rFonts")
    fonts.set(qn("w:asciiTheme"), "minorHAnsi")
    fonts.set(qn("w:hAnsiTheme"), "minorHAnsi")
    fonts.set(qn("w:cstheme"),    "minorHAnsi")
    if bold:
        etree.SubElement(rpr, qn("w:b"))
        etree.SubElement(rpr, qn("w:bCs"))
    if italic:
        etree.SubElement(rpr, qn("w:i"))
        etree.SubElement(rpr, qn("w:iCs"))
    sz = etree.SubElement(rpr, qn("w:sz"))
    sz.set(qn("w:val"), str(size))
    szcs = etree.SubElement(rpr, qn("w:szCs"))
    szcs.set(qn("w:val"), str(size))
    t = etree.SubElement(r, W+"t")
    t.text = text
    if text.startswith(" ") or text.endswith(" "):
        t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    return r

def yeni_paragraf(text: str = "", bold=False, italic=False,
                  numId: str = None, ilvl: str = "0",
                  extra_rpr_tags: list = None) -> etree._Element:
    """<w:p> oluşturur."""
    p = etree.Element(W+"p")
    ppr = etree.SubElement(p, W+"pPr")

    if numId:
        numpr = etree.SubElement(ppr, W+"numPr")
        ilvl_el = etree.SubElement(numpr, W+"ilvl")
        ilvl_el.set(qn("w:val"), ilvl)
        num_el = etree.SubElement(numpr, W+"numId")
        num_el.set(qn("w:val"), numId)

    rpr_in_ppr = etree.SubElement(ppr, W+"rPr")
    fonts = etree.SubElement(rpr_in_ppr, W+"rFonts")
    fonts.set(qn("w:asciiTheme"), "minorHAnsi")
    fonts.set(qn("w:hAnsiTheme"), "minorHAnsi")
    fonts.set(qn("w:cstheme"),    "minorHAnsi")
    sz = etree.SubElement(rpr_in_ppr, W+"sz"); sz.set(qn("w:val"), "22")
    etree.SubElement(rpr_in_ppr, W+"szCs").set(qn("w:val"), "22")

    if text:
        p.append(yeni_run(text, bold=bold, italic=italic))
    return p

def bos_paragraf() -> etree._Element:
    return yeni_paragraf()

def degerlendirme_tablosu_xml() -> list:
    """Değerlendirme tablosunu XML element listesi olarak döndürür."""
    # Şablondaki değerlendirme tablosu yapısını taklit et
    tbl = etree.Element(W+"tbl")

    tblPr = etree.SubElement(tbl, W+"tblPr")
    style = etree.SubElement(tblPr, W+"tblStyle"); style.set(qn("w:val"), "TabloKlavuzu")
    w_el  = etree.SubElement(tblPr, W+"tblW"); w_el.set(qn("w:w"), "0"); w_el.set(qn("w:type"), "auto")

    tblGrid = etree.SubElement(tbl, W+"tblGrid")
    gc1 = etree.SubElement(tblGrid, W+"gridCol"); gc1.set(qn("w:w"), "1575")
    gc2 = etree.SubElement(tblGrid, W+"gridCol"); gc2.set(qn("w:w"), "7487")

    # Satır 1: "Değerlendirme" etiketi
    tr1 = etree.SubElement(tbl, W+"tr")
    tc1 = etree.SubElement(tr1, W+"tc")
    tcPr1 = etree.SubElement(tc1, W+"tcPr")
    tcW1 = etree.SubElement(tcPr1, W+"tcW"); tcW1.set(qn("w:w"), "1575"); tcW1.set(qn("w:type"), "dxa")
    p1 = etree.SubElement(tc1, W+"p")
    ppr1 = etree.SubElement(p1, W+"pPr")
    rpr1 = etree.SubElement(ppr1, W+"rPr")
    b1 = etree.SubElement(rpr1, W+"b"); etree.SubElement(rpr1, W+"bCs")
    etree.SubElement(rpr1, W+"i"); etree.SubElement(rpr1, W+"iCs")
    sz1 = etree.SubElement(rpr1, W+"sz"); sz1.set(qn("w:val"), "22")
    etree.SubElement(rpr1, W+"szCs").set(qn("w:val"), "22")
    p1.append(yeni_run("Değerlendirme", bold=True, italic=True))

    # Satır 2: boş alan
    tr2 = etree.SubElement(tbl, W+"tr")
    tc2 = etree.SubElement(tr2, W+"tc")
    tcPr2 = etree.SubElement(tc2, W+"tcPr")
    tcW2 = etree.SubElement(tcPr2, W+"tcW"); tcW2.set(qn("w:w"), "1575"); tcW2.set(qn("w:type"), "dxa")
    sp2 = etree.SubElement(tcPr2, W+"vMerge")
    tc2.append(bos_paragraf())
    tc3 = etree.SubElement(tr2, W+"tc")
    tcPr3 = etree.SubElement(tc3, W+"tcPr")
    tcW3 = etree.SubElement(tcPr3, W+"tcW"); tcW3.set(qn("w:w"), "7487"); tcW3.set(qn("w:type"), "dxa")
    tc3.append(bos_paragraf())

    return [bos_paragraf(), tbl, bos_paragraf()]


def resim_xml(doc: Document, img_bytes: bytes, width_cm: float = 7.0) -> etree._Element:
    """Resmi belgeye ekler, <w:p> içinde <w:drawing> döndürür."""
    from docx.shared import Cm
    img_stream = BytesIO(img_bytes)
    try:
        pic_run = doc.add_picture(img_stream, width=Cm(width_cm))
        # Son eklenen paragrafı al ve body'den çıkar, XML döndür
        body = doc.element.body
        last_p = body[-2]  # add_picture son paragrafa ekler
        body.remove(last_p)
        return last_p
    except Exception as e:
        log.warning(f"Resim eklenemedi: {e}")
        return None


# ──────────────────────────────────────────────
#  WORD OLUŞTURMA
# ──────────────────────────────────────────────
def build_docx(s: dict) -> str:
    if not Path(TASLAK_PATH).exists():
        raise FileNotFoundError(f"Taslak bulunamadı: {TASLAK_PATH}")

    firma    = (s["firma"] or "rapor").replace(" ", "_")
    tarih_fn = s["tarih"].replace("/", "").replace(".", "")
    out_path = OUTPUT_DIR / f"Rapor_{firma}_{tarih_fn}.docx"
    shutil.copy2(TASLAK_PATH, out_path)

    doc  = Document(str(out_path))
    body = doc.element.body

    # ── 1. BAŞLIK TABLOLARINA VERİ YAZ ──
    # Tüm tabloları bul, sırayla doldur
    all_tables = doc.tables

    def tablo_hucre_yaz(tbl, row_idx: int, col_idx: int, text: str):
        try:
            cell = tbl.rows[row_idx].cells[col_idx]
            p = cell.paragraphs[0]
            # Mevcut run'ları temizle
            for run in p.runs:
                run.text = ""
            if p.runs:
                p.runs[0].text = text
            else:
                p.add_run(text)
        except:
            pass

    # Tablo 0: Tarih tablosu (1 satır, sağ hücre = tarih)
    if len(all_tables) > 0:
        tablo_hucre_yaz(all_tables[0], 0, 1, s["tarih"])

    # Tablo 1: İş emri, firma, ülke, satış (4 satır)
    if len(all_tables) > 1:
        t = all_tables[1]
        vals = [s["is_emri"], s["firma"], s["ulke_sehir"], s["satis"]]
        for i, v in enumerate(vals):
            if i < len(t.rows):
                tablo_hucre_yaz(t, i, 1, v)

    # Tablo 2: Montaj formeni
    if len(all_tables) > 2:
        tablo_hucre_yaz(all_tables[2], 0, 1, s["foremen"])

    # Tablo 3: Personel - her isim ayrı satıra (satır 1-5, sol sütun)
    if len(all_tables) > 3:
        t3 = all_tables[3]
        for pi, isim in enumerate(s["personel"][:5]):  # max 5 kişi
            row_idx = pi + 1  # 0. satır başlık, 1'den başla
            if row_idx < len(t3.rows):
                tablo_hucre_yaz(t3, row_idx, 0, isim)

    # Tablo 4: Başlangıç/Bitiş tarihleri
    if len(all_tables) > 4:
        t = all_tables[4]
        if len(t.rows) > 0 and len(t.rows[0].cells) > 1:
            tablo_hucre_yaz(t, 0, 1, s["baslangic"])
            tablo_hucre_yaz(t, 0, 3, s["bitis"])

    # Tablo 5: Adam gün / bitirme oranı (satır 1=gerçekleşen, satır 2=planlanan)
    if len(all_tables) > 5:
        t = all_tables[5]
        tablo_hucre_yaz(t, 1, 2, s["gerceklesen"])
        tablo_hucre_yaz(t, 2, 2, s["planlanan"])

    # Tablo 6: Tesis tipi — başlık satır 0 col 0, içerik satır 1 col 0
    if len(all_tables) > 6:
        tablo_hucre_yaz(all_tables[6], 1, 0, s["tesis_tipi"])

    # Tablo 7: Notlar
    if len(all_tables) > 7:
        tablo_hucre_yaz(all_tables[7], 0, 1, s["notlar"])

    # Tablo 8: İmza
    if len(all_tables) > 8:
        t = all_tables[8]
        tablo_hucre_yaz(t, 0, 0, s["foremen"])
        tablo_hucre_yaz(t, 0, 1, s["kontrol_eden"])
        tablo_hucre_yaz(t, 0, 2, s["onaylayan"])

    # ── 2. GÜNLÜK KURULUM RAPORU BÖLÜMÜNÜ DOLDUR ──
    # "Günlük Kurulum Raporu" paragrafını bul
    gunluk_p = None
    aksalik_p = None
    musteri_p = None
    oneri_p = None

    for p in body.iter(W+"p"):
        texts = "".join(t.text or "" for t in p.iter(W+"t"))
        if "Günlük Kurulum Raporu" in texts:
            gunluk_p = p
        elif "Kurulum esnasında karşılaşılan" in texts:
            aksalik_p = p
        elif "Müşteri kaynaklı sıkıntılar" in texts:
            musteri_p = p
        elif "Öneriler" in texts:
            oneri_p = p

    def paragraf_sonrasina_ekle(anchor_p, new_elements: list):
        """anchor_p'den sonra elemanları sırayla ekler."""
        parent = anchor_p.getparent()
        idx    = list(parent).index(anchor_p)
        for i, el in enumerate(new_elements):
            parent.insert(idx + 1 + i, el)

    def siradaki_bos_liste_paragraflarini_sil(anchor_p, numId_val):
        """Şablondaki boş liste paragraflarını temizler."""
        parent  = anchor_p.getparent()
        to_remove = []
        found   = False
        for el in parent:
            if el is anchor_p:
                found = True
                continue
            if not found:
                continue
            texts = "".join(t.text or "" for t in el.iter(W+"t"))
            numIds = [n.get(qn("w:val")) for n in el.iter(W+"numId")]
            if numId_val in numIds and not texts.strip():
                to_remove.append(el)
            elif texts.strip() in ("Çalışma saatleri", ":", "Fazla mesai", ""):
                # boş çalışma saati satırları
                all_t = [t.text or "" for t in el.iter(W+"t")]
                combined = "".join(all_t)
                if "Çalışma saatleri" in combined or "Fazla mesai" in combined or not combined.strip():
                    to_remove.append(el)
            else:
                break
        for el in to_remove:
            parent.remove(el)

    # ── Günlük kayıtlar (GRUPLANMIŞ VERSİYON) ──
    if gunluk_p is not None and s["gunluk"]:
        # Şablondaki boş liste + çalışma saati satırlarını temizle
        parent = gunluk_p.getparent()
        sil_listesi = []
        gec = False
        dur_metinler = {"Kurulum esnasında", "Müşteri kaynaklı", "Öneriler"}
        for el in parent:
            if el is gunluk_p:
                gec = True
                continue
            if not gec:
                continue
            metinler = "".join(t.text or "" for t in el.iter(W+"t"))
            if any(m in metinler for m in dur_metinler):
                break
            sil_listesi.append(el)
        for el in sil_listesi:
            parent.remove(el)

        # AYNI TARİHTEKİ KAYITLARI GRUPLA
        tarih_gruplari = {}
        for g in s["gunluk"]:
            tarih = g.get("tarih", "")
            if tarih not in tarih_gruplari:
                tarih_gruplari[tarih] = {
                    "tarih": tarih,
                    "gun": g.get("gun", gun_adi(tarih)),
                    "maddeler": [],
                    "calisma_saati": g.get("calisma_saati", ""),
                    "fazla_mesai": g.get("fazla_mesai", "")
                }
            # Maddeleri birleştir
            tarih_gruplari[tarih]["maddeler"].extend(g.get("maddeler", []))
            
            # Çalışma saatleri ve fazla mesai bilgilerini güncelle (son gelen geçerli olsun)
            if g.get("calisma_saati"):
                tarih_gruplari[tarih]["calisma_saati"] = g.get("calisma_saati")
            if g.get("fazla_mesai"):
                tarih_gruplari[tarih]["fazla_mesai"] = g.get("fazla_mesai")

        # Yeni günlük kayıtları ekle (gruplanmış)
        yeni_els = []
        yeni_els.append(bos_paragraf())
        
        # Tarihe göre sırala (eskiden yeniye) — sonradan girilen geçmiş tarihler doğru yere girer
        sirali_tarihler = sorted(
            tarih_gruplari.items(),
            key=lambda x: datetime.strptime(x[0], "%d.%m.%Y") if x[0] else datetime.min
        )

        # Gruplanmış kayıtları sıralı olarak ekle
        for i, (tarih, grup) in enumerate(sirali_tarihler, 1):
            gun = grup.get("gun", gun_adi(tarih))
            baslik = f"{i}.  {tarih} {gun}".strip()

            # Düz numara — numId yok, Word listesiyle çakışmasın
            p_num = yeni_paragraf(baslik, bold=True)
            yeni_els.append(p_num)

            # Maddeleri ekle (her madde ayrı satırda)
            for madde in grup.get("maddeler", []):
                yeni_els.append(yeni_paragraf(f"- {madde}"))

            # Çalışma saatleri
            p_saat = etree.Element(W+"p")
            p_saat.append(yeni_run("Çalışma saatleri : ", bold=False))
            p_saat.append(yeni_run(grup.get("calisma_saati", ""), bold=False))
            yeni_els.append(p_saat)

            # Fazla mesai
            p_fazla = etree.Element(W+"p")
            p_fazla.append(yeni_run("Fazla mesai : ", bold=False))
            p_fazla.append(yeni_run(grup.get("fazla_mesai", ""), bold=False))
            yeni_els.append(p_fazla)
            yeni_els.append(bos_paragraf())

        paragraf_sonrasina_ekle(gunluk_p, yeni_els)

    # ── Aksaklıklar ──
    if aksalik_p is not None and s["aksalik"]:
        parent = aksalik_p.getparent()
        sil_listesi = []
        gec = False
        dur_metinler = {"Müşteri kaynaklı", "Öneriler"}
        for el in parent:
            if el is aksalik_p:
                gec = True
                continue
            if not gec:
                continue
            metinler = "".join(t.text or "" for t in el.iter(W+"t"))
            if any(m in metinler for m in dur_metinler):
                break
            # Tablo veya boş paragraf → sil
            sil_listesi.append(el)
        for el in sil_listesi:
            parent.remove(el)

        yeni_els = [bos_paragraf()]
        for i, a in enumerate(s["aksalik"], 1):
            # Düz numara — numId yok
            p_num = yeni_paragraf(f"{i}.  {a['aciklama']}")
            yeni_els.append(p_num)

            # Fotoğraflar
            for foto_bytes in a.get("fotograflar", []):
                try:
                    img_stream = BytesIO(foto_bytes)
                    from docx.shared import Cm as DocxCm
                    p_img = etree.Element(W+"p")
                    # Fotoğrafı geçici olarak doc'a ekle, drawing XML'i al
                    tmp_para = doc.add_paragraph()
                    run = tmp_para.add_run()
                    run.add_picture(img_stream, width=DocxCm(7))
                    drawing = tmp_para._p.find(".//" + qn("w:drawing"))
                    if drawing is not None:
                        p_img.append(copy.deepcopy(drawing))
                    # Geçici paragrafı kaldır
                    tmp_para._p.getparent().remove(tmp_para._p)
                    yeni_els.append(p_img)
                except Exception as e:
                    log.warning(f"Foto eklenemedi: {e}")

            # Değerlendirme tablosu
            yeni_els.extend(degerlendirme_tablosu_xml())

        paragraf_sonrasina_ekle(aksalik_p, yeni_els)

    # ── Müşteri sıkıntıları ──
    if musteri_p is not None and s["musteri"]:
        parent = musteri_p.getparent()
        sil_listesi = []
        gec = False
        for el in parent:
            if el is musteri_p:
                gec = True
                continue
            if not gec:
                continue
            metinler = "".join(t.text or "" for t in el.iter(W+"t"))
            if "Öneriler" in metinler:
                break
            sil_listesi.append(el)
        for el in sil_listesi:
            parent.remove(el)

        yeni_els = [bos_paragraf()]
        for i, m in enumerate(s["musteri"], 1):
            yeni_els.append(yeni_paragraf(f"{i}.  {m}"))
        paragraf_sonrasina_ekle(musteri_p, yeni_els)

    # ── Öneriler ──
    if oneri_p is not None and s["oneriler"]:
        parent = oneri_p.getparent()
        # Sonrasındaki boş maddeleri temizle
        sil_listesi = []
        gec = False
        for el in parent:
            if el is oneri_p:
                gec = True
                continue
            if not gec:
                continue
            if el.tag == W+"sectPr":
                break
            sil_listesi.append(el)
        for el in sil_listesi:
            parent.remove(el)

        yeni_els = [bos_paragraf()]
        for i, o in enumerate(s["oneriler"], 1):
            yeni_els.append(yeni_paragraf(f"{i}.  {o}"))
        paragraf_sonrasina_ekle(oneri_p, yeni_els)

    doc.save(str(out_path))
    return str(out_path)


# ──────────────────────────────────────────────
#  BOT MENÜ & KOMUTLAR
# ──────────────────────────────────────────────
MOD_LABELS = {
    "gunluk": "📅 Günlük Kayıt",
    "toplu":  "📦 Toplu Giriş",
    "aksalik": "⚠️ Aksaklık / Sorun",
    "musteri": "👷 Müşteri Sıkıntısı",
    "oneri":   "💡 Öneri",
}

def ana_menu():
    return InlineKeyboardMarkup([
        [InlineKeyboardButton("📅 Günlük kayıt ekle",     callback_data="mod_gunluk")],
        [InlineKeyboardButton("📦 Toplu giriş (tarih aralığı)", callback_data="mod_toplu")],
        [InlineKeyboardButton("⚠️ Aksaklık / sorun ekle", callback_data="mod_aksalik")],
        [InlineKeyboardButton("👷 Müşteri sıkıntısı ekle", callback_data="mod_musteri")],
        [InlineKeyboardButton("💡 Öneri ekle",             callback_data="mod_oneri")],
        [InlineKeyboardButton("📄 Word dosyası al",        callback_data="kaydet")],
        [InlineKeyboardButton("📋 Özet göster",            callback_data="ozet")],
    ])

def onay_kb():
    return InlineKeyboardMarkup([[
        InlineKeyboardButton("✅ Onayla & Ekle", callback_data="confirm"),
        InlineKeyboardButton("❌ İptal",          callback_data="cancel"),
    ]])

async def cmd_start(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    sessions.pop(uid, None); get_session(uid)
    await update.message.reply_text(
        "👷 *Montaj Rapor Botuna Hoş Geldiniz!*\n\n"
        "1️⃣ Önce /ayarla ile proje bilgilerini girin\n"
        "2️⃣ Menüden bölüm seçin ve sesli/yazılı/fotoğraf gönderin\n"
        "3️⃣ /kaydet ile Word dosyasını alın\n\n"
        "⚠️ `MONTAJ_RAPORU_TASLAK.docx` dosyasının bot ile aynı klasörde olması gerekiyor!",
        parse_mode="Markdown", reply_markup=ana_menu())

async def cmd_ayarla(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    get_session(update.effective_user.id)["_mod"] = "setup"
    await update.message.reply_text(
        "Proje bilgilerini şu formatta gönderin:\n\n"
        "`IS_EMRI: 2025025\n"
        "FIRMA: Palkana\n"
        "ULKE_SEHIR: Irak/Süleymaniye\n"
        "SATIS: Mehmet SEVGİLİ\n"
        "FOREMEN: Tolga ÇELİK\n"
        "PERSONEL: Atıf USLU (Konstrüksiyon), Murat UYSAL (Kabin)\n"
        "BASLANGIC: 28.11.2025\n"
        "BITIS: 25.12.2025\n"
        "GERCEKLESEN: 110\n"
        "PLANLANAN: 100\n"
        "TESIS: 4 Banyolu Püskürtme Yıkamalı Tesis\n"
        "NOTLAR: Tesis devreye alınmamıştır.\n"
        "KONTROL: Emre GÜLERYÜZ\n"
        "ONAYLAYAN: İzzet DAĞBAŞI`",
        parse_mode="Markdown")

async def cmd_kaydet(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    s = get_session(update.effective_user.id)
    await update.message.reply_text(
        f"⚠️ *Montaj tamamlandı mı?*\n\n"
        f"Firma: {s['firma'] or '—'}\n"
        f"Günlük kayıt: {len(s['gunluk'])} gün  |  Aksaklık: {len(s['aksalik'])}\n\n"
        "Raporu Word olarak almak istediğinizden emin misiniz?",
        parse_mode="Markdown",
        reply_markup=InlineKeyboardMarkup([[
            InlineKeyboardButton("✅ Evet, dosyayı al", callback_data="kaydet_onayla"),
            InlineKeyboardButton("❌ Hayır, devam et", callback_data="kaydet_iptal"),
        ]])
    )

async def cmd_ozet(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    s = get_session(update.effective_user.id)
    lines = [f"📋 *{s['tarih']}*  |  {s['firma'] or '—'}",
             f"Günlük: {len(s['gunluk'])} gün  |  Aksaklık: {len(s['aksalik'])}",
             f"Müşteri: {len(s['musteri'])}  |  Öneri: {len(s['oneriler'])}"]
    await update.message.reply_text("\n".join(lines), parse_mode="Markdown", reply_markup=ana_menu())

async def cmd_sifirla(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    sessions.pop(update.effective_user.id, None); get_session(update.effective_user.id)
    await update.message.reply_text("🔄 Yeni rapor başlatıldı!", reply_markup=ana_menu())


# ──────────────────────────────────────────────
#  MESAJ İŞLEYİCİLERİ
# ──────────────────────────────────────────────
async def isle_metin(uid: int, text: str, update: Update):
    s   = get_session(uid)
    mod = s.get("_mod")

    # ── Setup ──
    if mod == "setup":
        for line in text.splitlines():
            u = line.upper()
            v = line.split(":", 1)[1].strip() if ":" in line else ""
            if   "IS_EMRI:"   in u: s["is_emri"]   = v
            elif "FIRMA:"     in u: s["firma"]      = v
            elif "ULKE_SEHIR:"in u: s["ulke_sehir"] = v
            elif "SATIS:"     in u: s["satis"]      = v
            elif "FOREMEN:"   in u: s["foremen"]    = v
            elif "PERSONEL:"  in u: s["personel"]   = [p.strip() for p in v.split(",")]
            elif "BASLANGIC:" in u: s["baslangic"]  = v
            elif "BITIS:"     in u: s["bitis"]      = v
            elif "GERCEKLESEN:"in u: s["gerceklesen"]= v
            elif "PLANLANAN:" in u: s["planlanan"]  = v
            elif "TESIS:"     in u: s["tesis_tipi"] = v
            elif "NOTLAR:"    in u: s["notlar"]     = v
            elif "KONTROL:"   in u: s["kontrol_eden"]= v
            elif "ONAYLAYAN:" in u: s["onaylayan"]  = v
        s["_mod"] = None
        await update.message.reply_text(
            f"✅ Kaydedildi! Firma: {s['firma']}  |  Foremen: {s['foremen']}",
            reply_markup=ana_menu())
        return

    if mod is None:
        await update.message.reply_text("Lütfen önce ne eklemek istediğinizi seçin:", reply_markup=ana_menu())
        return

    # ── Aksalık açıklama modu (fotoğraf sonrası metin) ──
    if mod == "aksalik_metin":
        thinking = await update.message.reply_text("⏳ GPT-4o işliyor...")
        fmt = await gpt_aksalik(text)
        s["_pending_fmt"]  = fmt
        s["_pending_text"] = text
        await thinking.edit_text(
            f"⚠️ *Aksaklık:*\n\n{fmt}\n\n"
            f"📸 {len(s['_pending_photos'])} fotoğraf eklendi.\n"
            "Daha fazla fotoğraf gönderin veya onaylayın:",
            reply_markup=onay_kb(), parse_mode="Markdown")
        s["_mod"] = "aksalik_foto"
        return

    thinking = await update.message.reply_text("⏳ GPT-4o işliyor...")
    try:
        if mod == "toplu":
            # Tarih aralığını metinden çıkar, geri kalanı paragraf olarak gönder
            # Format beklentisi: "GG.AA.YYYY - GG.AA.YYYY ... açıklama"
            import re as _re
            tarih_pattern = _re.search(r"(\d{2}\.\d{2}\.\d{4})\s*[-–]\s*(\d{2}\.\d{2}\.\d{4})", text)
            if tarih_pattern:
                aralik = f"{tarih_pattern.group(1)} - {tarih_pattern.group(2)}"
                paragraf = text[tarih_pattern.end():].strip() or text
            else:
                aralik = f"{bugun()} - {bugun()}"
                paragraf = text

            gunler = await gpt_toplu(aralik, paragraf)
            ozet = f"📦 *Toplu giriş — {aralik}*\n_{len(gunler)} güne bölündü:_\n\n"
            for g in gunler[:5]:  # önizlemede max 5 gün göster
                ozet += f"📅 *{g.get('tarih','')} {g.get('gun','')}*\n"
                for m in g.get("maddeler",[])[:2]:
                    ozet += f"- {m}\n"
                ozet += "\n"
            if len(gunler) > 5:
                ozet += f"_...ve {len(gunler)-5} gün daha_\n"
            s["_pending_fmt"]  = gunler
            s["_pending_text"] = text
            await thinking.edit_text(
                ozet + "\nOnaylıyor musunuz? (Tüm günler rapora eklenecek)",
                reply_markup=onay_kb(), parse_mode="Markdown")

        elif mod == "gunluk":
            data = await gpt_gunluk(text)
            gun  = data.get("gun", "")
            ozet = f"📅 *{data.get('tarih','')} {gun}*\n"
            for m in data.get("maddeler", []):
                ozet += f"- {m}\n"
            ozet += f"\nÇalışma: {data.get('calisma_saati','')}  |  Fazla: {data.get('fazla_mesai','—') or '—'}"
            s["_pending_fmt"]  = data
            s["_pending_text"] = text
            await thinking.edit_text(ozet + "\n\nOnaylıyor musunuz?", reply_markup=onay_kb(), parse_mode="Markdown")

        elif mod == "aksalik":
            fmt = await gpt_aksalik(text)
            s["_pending_fmt"]  = fmt
            s["_pending_text"] = text
            await thinking.edit_text(
                f"⚠️ *Aksaklık:*\n\n{fmt}\n\n"
                "📸 Fotoğraf eklemek ister misiniz? Şimdi fotoğraf gönderin veya direkt onaylayın:",
                reply_markup=onay_kb(), parse_mode="Markdown")
            s["_mod"] = "aksalik_foto"

        elif mod in ("musteri", "oneri"):
            maddeler = await gpt_liste(text)
            ozet = "\n".join(f"• {m}" for m in maddeler)
            s["_pending_fmt"]  = maddeler
            s["_pending_text"] = text
            lbl = "👷 Müşteri sıkıntısı" if mod == "musteri" else "💡 Öneriler"
            await thinking.edit_text(f"{lbl}:\n\n{ozet}\n\nOnaylıyor musunuz?", reply_markup=onay_kb(), parse_mode="Markdown")

    except Exception as e:
        log.exception("GPT hatası")
        await thinking.edit_text(f"❌ GPT Hata: {e}")

async def handle_text(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    await isle_metin(update.effective_user.id, update.message.text.strip(), update)

async def handle_voice(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    s   = get_session(uid)
    if s.get("_mod") is None:
        await update.message.reply_text("Lütfen önce ne eklemek istediğinizi seçin:", reply_markup=ana_menu())
        return
    status = await update.message.reply_text("🎤 Whisper transkript ediyor...")
    try:
        vf = await update.message.voice.get_file()
        with tempfile.NamedTemporaryFile(suffix=".ogg", delete=False) as tmp:
            await vf.download_to_drive(tmp.name); tmp_path = tmp.name
        transcript = await transcribe(tmp_path)
        os.unlink(tmp_path)
        await status.edit_text(f"🎤 _{transcript}_\n\n⏳ İşleniyor...", parse_mode="Markdown")
        # Doğrudan transcript metnini gönder
        await isle_metin(uid, transcript, update)
    except Exception as e:
        log.exception("Ses hatası")
        await status.edit_text(f"❌ Ses Hata: {e}")

async def handle_photo(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    uid = update.effective_user.id
    s   = get_session(uid)
    mod = s.get("_mod")

    if mod not in ("aksalik", "aksalik_foto", "aksalik_metin"):
        await update.message.reply_text("Fotoğraf eklemek için ⚠️ Aksaklık modunu seçin.", reply_markup=ana_menu())
        return

    photo = update.message.photo[-1]
    f     = await photo.get_file()
    buf   = BytesIO()
    await f.download_to_memory(buf)
    s["_pending_photos"].append(buf.getvalue())
    s["_mod"] = "aksalik_foto"

    caption = (update.message.caption or "").strip()
    if caption:
        # Başlıkla açıklama geldi → direkt işle
        thinking = await update.message.reply_text("⏳ İşleniyor...")
        fmt = await gpt_aksalik(caption)
        s["_pending_fmt"]  = fmt
        s["_pending_text"] = caption
        await thinking.edit_text(
            f"⚠️ *Aksaklık:*\n\n{fmt}\n\n📸 {len(s['_pending_photos'])} fotoğraf.\n\nOnaylıyor musunuz?",
            reply_markup=onay_kb(), parse_mode="Markdown")
    else:
        kb = onay_kb() if s.get("_pending_fmt") else None
        await update.message.reply_text(
            f"📸 {len(s['_pending_photos'])} fotoğraf alındı.\n"
            "Açıklamayı yazın/sesli gönderin" +
            (" veya doğrudan onaylayın:" if s.get("_pending_fmt") else ":"),
            reply_markup=kb)
        if not s.get("_pending_fmt"):
            s["_mod"] = "aksalik_metin"


# ──────────────────────────────────────────────
#  CALLBACK HANDLER
# ──────────────────────────────────────────────
async def handle_callback(update: Update, ctx: ContextTypes.DEFAULT_TYPE):
    query = update.callback_query
    uid   = query.from_user.id
    s     = get_session(uid)
    await query.answer()
    data  = query.data

    if data.startswith("mod_"):
        mod = data.replace("mod_", "")
        s["_mod"] = mod
        s["_pending_photos"] = []
        s["_pending_fmt"]    = None
        s["_pending_text"]   = ""
        msgs = {
            "gunluk":  "📅 Günlük çalışmayı anlatın (sesli veya yazılı).\nÖrnek: *'Bugün yıkama hattı ring boruları tamamlandı, konveyör montajı yapıldı'*",
            "toplu":   "📦 *Toplu giriş*\n\nTarih aralığını ve o dönemde yapılan tüm işleri tek paragraf olarak yazın.\n\nÖrnek:\n_02.02.2026 - 13.03.2026 tarihleri arasında yıkama hattı kuruldu, konveyör montajı yapıldı, elektrik bağlantıları tamamlandı, devreye alma testleri yapıldı._\n\nGPT-4o gün gün bölerek rapora ekleyecek.",
            "aksalik": "⚠️ Fotoğraf gönderin (alt yazı olarak açıklama ekleyin)\nveya yazılı/sesli açıklayın.",
            "musteri": "👷 Müşteri kaynaklı sıkıntıları anlatın:",
            "oneri":   "💡 Önerilerinizi anlatın:",
        }
        await query.edit_message_text(f"*{MOD_LABELS.get(mod, mod)}*\n\n{msgs.get(mod, '')}", parse_mode="Markdown")
        return

    if data == "kaydet":
        # Onay iste
        await query.edit_message_text(
            f"⚠️ *Montaj tamamlandı mı?*\n\n"
            f"Firma: {s['firma'] or '—'}\n"
            f"Günlük kayıt: {len(s['gunluk'])} gün  |  Aksaklık: {len(s['aksalik'])}\n\n"
            "Raporu Word olarak almak istediğinizden emin misiniz?",
            parse_mode="Markdown",
            reply_markup=InlineKeyboardMarkup([[
                InlineKeyboardButton("✅ Evet, dosyayı al", callback_data="kaydet_onayla"),
                InlineKeyboardButton("❌ Hayır, devam et", callback_data="kaydet_iptal"),
            ]])
        )
        return

    if data == "kaydet_onayla":
        await query.edit_message_text("📄 Word hazırlanıyor...")
        try:
            path = build_docx(s)
            with open(path, "rb") as f:
                await query.message.reply_document(
                    document=f, filename=Path(path).name,
                    caption=f"✅ *{s['firma'] or 'Rapor'} — {s['tarih']}*",
                    parse_mode="Markdown")
            await query.message.reply_text("Rapor teslim edildi. Yeni rapor için /sifirla", reply_markup=ana_menu())
        except Exception as e:
            await query.message.reply_text(f"❌ Hata: {e}")
        return

    if data == "kaydet_iptal":
        await query.edit_message_text("↩️ Rapora devam ediliyor.", reply_markup=ana_menu())
        return

    if data == "ozet":
        lines = [f"📋 *{s['tarih']}*  |  {s['firma'] or '—'}",
                 f"Günlük: {len(s['gunluk'])} gün  |  Aksaklık: {len(s['aksalik'])}",
                 f"Müşteri: {len(s['musteri'])}  |  Öneri: {len(s['oneriler'])}"]
        await query.edit_message_text("\n".join(lines), parse_mode="Markdown", reply_markup=ana_menu())
        return

    if data == "confirm":
        mod = s.get("_mod") or ""
        fmt = s.get("_pending_fmt")

        if mod == "toplu" and isinstance(fmt, list) and fmt and isinstance(fmt[0], dict):
            s["gunluk"].extend(fmt)
            # Tarihe göre sırala
            from datetime import datetime as _dt
            s["gunluk"].sort(key=lambda x: _dt.strptime(x.get("tarih", "01.01.2000"), "%d.%m.%Y"))
            await query.edit_message_text(
                f"✅ Toplu giriş eklendi! {len(fmt)} gün rapora eklendi. (Toplam {len(s['gunluk'])} gün)",
                reply_markup=ana_menu())

        elif "gunluk" in mod and isinstance(fmt, dict):
            s["gunluk"].append(fmt)
            await query.edit_message_text(f"✅ Günlük eklendi! ({len(s['gunluk'])} gün)", reply_markup=ana_menu())

        elif "aksalik" in mod and fmt:
            s["aksalik"].append({"aciklama": fmt, "fotograflar": list(s["_pending_photos"])})
            s["_pending_photos"] = []
            await query.edit_message_text(
                f"✅ Aksaklık eklendi! ({len(s['aksalik'])}. — {len(s['aksalik'][-1]['fotograflar'])} fotoğraf)",
                reply_markup=ana_menu())

        elif mod == "musteri" and isinstance(fmt, list):
            s["musteri"].extend(fmt)
            await query.edit_message_text(f"✅ {len(fmt)} müşteri sıkıntısı eklendi!", reply_markup=ana_menu())

        elif mod == "oneri" and isinstance(fmt, list):
            s["oneriler"].extend(fmt)
            await query.edit_message_text(f"✅ {len(fmt)} öneri eklendi!", reply_markup=ana_menu())

        else:
            await query.edit_message_text("⚠️ Onaylanacak veri yok.", reply_markup=ana_menu())

        s["_mod"] = None; s["_pending_fmt"] = None; s["_pending_text"] = ""
        return

    if data == "cancel":
        s["_mod"] = None; s["_pending_fmt"] = None
        s["_pending_text"] = ""; s["_pending_photos"] = []
        await query.edit_message_text("❌ İptal edildi.", reply_markup=ana_menu())


# ──────────────────────────────────────────────
#  UYGULAMA
# ──────────────────────────────────────────────
def main():
    app = Application.builder().token(TELEGRAM_TOKEN).build()
    app.add_handler(CommandHandler("start",   cmd_start))
    app.add_handler(CommandHandler("ayarla",  cmd_ayarla))
    app.add_handler(CommandHandler("kaydet",  cmd_kaydet))
    app.add_handler(CommandHandler("ozet",    cmd_ozet))
    app.add_handler(CommandHandler("sifirla", cmd_sifirla))
    app.add_handler(MessageHandler(filters.PHOTO,                   handle_photo))
    app.add_handler(MessageHandler(filters.VOICE,                   handle_voice))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_text))
    app.add_handler(CallbackQueryHandler(handle_callback))
    log.info("✅ Bot başlatıldı")
    app.run_polling(drop_pending_updates=True)

if __name__ == "__main__":
    main()